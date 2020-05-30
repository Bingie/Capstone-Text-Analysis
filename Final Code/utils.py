from docx import Document         
import re
import string
import json
import matplotlib.pyplot as plt
import seaborn as sns
from wordcloud import WordCloud
import numpy as np
import pandas as pd

#nltk.download('wordnet')
#nltk.download('stopwords')
#nltk.download('punkt')

# Delete the unnecessary marks in a word
def word_processor(word):
    new_word = ""
    for i in word:
        if ord(i)-ord('a') < 0 or ord(i)-ord('a') > 25:
            continue
        new_word = new_word + i
    return new_word


def clean_text(text, lemma, en_stop = [], exclude_sent = [], minwords = 2):
    """
    1. Transforms the text to lower case 
    2. removes punctuation and digits
    3. uses lemmatization to standarize words
    4. removes stopwords
    5. optional: removes sentences in exclude_sent
    
    comments with less words than minwords are discarded

    Parameters
    ----------
    text : TYPE: string
        DESCRIPTION: Comment of the survey 

    Returns
    -------
    text : TYPE: string 
        DESCRIPTION: clean version of the comment 

    """
    
    preprocessed_text = None
    
    text = str(text) #Some text is just numbers or empty
    text = text.lower() #lowercases every word 
    text = re.sub('[%s]'% re.escape(string.punctuation)," ",text) #removes punctuation
    text = re.sub('\w*\d\w','', text) #removes digits
    tokens = text.split()
    tokens = [word for word in tokens if word not in en_stop]
    
    if lemma:
        tokens = [lemma.lemmatize(word) for word in tokens]

    if len(tokens) >= minwords and text not in exclude_sent: 
        preprocessed_text = ' '.join(tokens)
    
    return preprocessed_text


def frequency_calculation(corpus, model, word1, word2, threshold = 0.5, diff = 0.1):
    
    # Two dicts to save the similar words with corresponding frequency
    word1_dict = {}
    word2_dict = {}
        
    # A dict that records the similarity of each word
    word_similarity_dict = {}
    discarded_words = []
    
    for doc in corpus:
        for word in doc:
            try:
                # word = word_processor(word)
                word1_sim = model.similarity(word, word1)
                word2_sim = model.similarity(word, word2)
                word_similarity_dict[word] = [word1_sim, word2_sim]
                
                if word1_sim - word2_sim > diff and word1_sim > threshold:
                    try:
                        word1_dict[word] += 1
                    except KeyError:
                        word1_dict[word] = 1
                
                if word2_sim - word1_sim > diff and word2_sim > threshold:
                    try:
                        word2_dict[word] += 1
                    except KeyError:
                        word2_dict[word] = 1
                    
            except KeyError:
               discarded_words.append(word) 
    
    return word1_dict, word2_dict, word_similarity_dict, discarded_words



def frequency_similarity_plot(words, frequency, similarity, topword):
    """
    

    Parameters
    ----------
    vec : TYPE
        DESCRIPTION.

    Returns
    -------
    None.

    """
    
    fig, ax1 = plt.subplots()
    ax1.set_xlabel("Words")
    ax1.set_ylabel("Frequency")
    ax1.bar(words, frequency, width=0.5, label='Frequency', color="blue")
    ax2 = ax1.twinx()
    ax2.set_ylim(0, 1)
    ax2.plot(words, similarity, c="red", linestyle='--', label="Similarity")
    ax2.set_ylabel("Cosine Similarity")
    plt.legend()
    plt.title("Frequency and similarity chart of words related to {}".format(topword))
    

def get_all_context(filename):
    doc = Document(filename)

    paragraph_text = []
    for paragraph in doc.paragraphs:
        # Sentences less than length of 5 is not part of the questionnaire
        words_split = str.split(paragraph.text)
        if len(words_split) > 5:
            paragraph_text.append(words_split)
    table_text = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                words_split = str.split(cell.text)
                words_split = [word_processor(word) for word in words_split]
                if len(words_split) > 5:
                    table_text.append(words_split)
    return paragraph_text, table_text


def get_all_words(paragraph_text, table_text, savepath):
    words_list = {}

    for comment in paragraph_text:
        for word in comment:
            word = word_processor(word)
            if word == "":
                continue
            try:
                words_list[word] += 1
            except:
                words_list[word] = 1

    for comment in table_text:
        for word in comment:
            word = word_processor(word)
            if word == "":
                continue
            try:
                words_list[word] += 1
            except:
                words_list[word] = 1
    with open(savepath, "w") as fout:
        json.dump(words_list, fout, indent=4)
    fout.close()
    

def WordCloud_generator(frequencies, title, **kwargs):
    
    width = 480
    height = 480
    max_words = 30
    
    
    if "width" in kwargs.keys():
        width = kwargs["width"]
        
    if "height" in kwargs.keys():
        height = kwargs["height"]
        
    if "max_words" in kwargs.keys():
        max_words = kwargs["max_words"]
    
    fig, ax = plt.subplots()
        
    wordcloud = WordCloud(width = width, height = height, max_words = max_words)
    wordcloud.fit_words(frequencies)
    
    ax.imshow(wordcloud, interpolation="bilinear")
    ax.axis("off")
    ax.set_title(title)
    ax.margins(x=0, y=0)
    fig.tight_layout(pad=0)
    

def tsnescatterplot(model, word, comparison_list = []):
    """ Plot in seaborn the results from the t-SNE dimensionality reduction 
    algorithm of the vectors of a query word, its list of most similar words, 
    and a list of words.
    """
    
    word_vector = model.wv.get_vector(word)
    ncols = np.shape(word_vector)[0]
    arrays = np.empty((0, ncols), dtype='f')
    word_vector = word_vector.reshape((1,ncols))
    arrays = np.append(arrays, word_vector, axis=0)
 
    word_labels = [word]
    color_list  = ['red']
    
    close_words = model.wv.most_similar(word)
    
    for wrd_score in close_words:
        wrd_vector = model.wv.get_vector(wrd_score[0])
        wrd_vector = wrd_vector.reshape((1,ncols))
        word_labels.append(wrd_score[0])
        color_list.append('blue')
        arrays = np.append(arrays, wrd_vector, axis=0)
    
    for word_compare in comparison_list:
        wrd_vector = model.wv.get_vector(word_compare)
        wrd_vector = wrd_vector.reshape((1,ncols))
        word_labels.append(word_compare)
        color_list.append('green')
        arrays = np.append(arrays, wrd_vector, axis=0)
   
    simplified_model = PCA(n_components = 0.9).fit_transform(arrays)
    
    # Finds t-SNE coordinates for 2 dimensions
    np.set_printoptions(suppress=True)
    
    tse = TSNE(n_components=2, 
               perplexity=40, 
               learning_rate = 600, 
               n_iter=5000, 
               method='exact')
    
    Y = tse.fit_transform(simplified_model)
    
     # Sets everything up to plot
    df = pd.DataFrame({'x': [x for x in Y[:, 0]],
                       'y': [y for y in Y[:, 1]],
                       'words': word_labels,
                       'color': color_list})
    
    fig, _ = plt.subplots()
    
    # Basic plot
    p1 = sns.regplot(data=df,
                     x="x",
                     y="y",
                     fit_reg=False,
                     marker="o",
                     scatter_kws={'s': 10,
                                  'facecolors': df['color']})
    
     # Adds annotations one by one with a loop
    for line in range(0, df.shape[0]):
         p1.text(df["x"][line],
                 df['y'][line],
                 '  ' + df["words"][line].title(),
                 horizontalalignment='left',
                 verticalalignment='bottom', size='medium',
                 color=df['color'][line],
                 weight='normal'
                ).set_size(15)

    
    plt.xlim(Y[:, 0].min()-50, Y[:, 0].max()+50)
    plt.ylim(Y[:, 1].min()-50, Y[:, 1].max()+50)
            
    plt.title('t-SNE visualization for {}'.format(word.title()))


def get_all_corpus():
    all_files = ["comp_1.docx", "comp_2.docx", "comp_3.docx", "comp_4.docx"]
    all_data = []
    for f in all_files:
        para, tble = get_all_context(f)
        all_data += para
        all_data += tble
    return all_data


def remove_negation_words(words):
    
    negation_words = ["no", "not", "nor"]
    for word in words:
        if word[-3:] == "n\'t" or word[-2:] == "n\'" or word in negation_words:
            words.remove(word)
    
    return words



    